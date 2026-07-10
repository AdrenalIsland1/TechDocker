"""Demo DOCX updater for the GitHub automation pipeline.

Run as ``python3 -m src.demo_docx_updater`` — inside GitHub Actions (where the
``GITHUB_*`` environment variables are provided) or locally with sensible
fallbacks (repository ``TechDocker``, current branch, ``HEAD~1..HEAD``).

Skeleton-aware flow (no full parse on routine pushes):

1. detect the changed files of the push,
2. load the stored skeleton JSON (build it first — one full parse — if absent),
3. build a simple change summary and route it via :mod:`src.change_router`,
4. insert a marked update block under the routed heading (found by scanning
   paragraph text, not by re-parsing), or append a new section,
5. update the skeleton JSON only when a new section was created.

If the routed heading cannot be found in the DOCX, the block is appended to
the end with a warning. SharePoint and LLM routing are future phases.
"""

from __future__ import annotations

import os
import subprocess
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Mapping, Optional

from docx import Document

from src.automation_demo import extract_repository_name, is_missing_sha
from src.change_router import CREATE_NEW, RoutingDecision, route_change
from src.document_skeleton_builder import (
    build_and_save_skeleton,
    skeleton_path_for,
)
from src.git_change_detector import ChangedFile, build_change_set
from src.project_resolver import ProjectConfig, resolve_project
from src.skeleton_store import (
    append_section,
    find_section_by_id,
    load_skeleton,
    save_skeleton,
)

AUTOMATED_SECTION_TITLE = "Automated Documentation Update"
AUTOMATED_SECTION_NOTE = (
    "This section was generated automatically by the TechDocker GitHub "
    "automation demo."
)


@dataclass
class UpdateResult:
    """What one updater run did, for printing and for tests."""

    document_path: Path
    project_id: str
    repository: str
    branch: str
    changed_files: list[ChangedFile] = field(default_factory=list)
    decision: Optional[RoutingDecision] = None
    placement: str = ""  # e.g. "under 'System Overview'" / "appended to end"
    skeleton_path: Optional[Path] = None
    skeleton_created: bool = False
    skeleton_updated: bool = False
    warnings: list[str] = field(default_factory=list)


def _local_git_output(args: list[str], repo_path: str) -> Optional[str]:
    """Run a git query for local fallbacks; ``None`` when it fails."""
    try:
        result = subprocess.run(
            ["git", *args],
            cwd=repo_path,
            capture_output=True,
            text=True,
            check=True,
        )
    except (subprocess.CalledProcessError, OSError):
        return None
    return result.stdout.strip() or None


def build_change_summary(changed_files: list[ChangedFile]) -> str:
    """One-line textual summary of a push's changed files."""
    if not changed_files:
        return "No changed files were available for this run."
    parts = [f"{changed.change_type} {changed.path}" for changed in changed_files]
    return f"{len(changed_files)} file(s) changed: " + "; ".join(parts)


# ---------------------------------------------------------------------------
# DOCX writing helpers
# ---------------------------------------------------------------------------
def _style_exists(document: Document, style_name: str) -> bool:
    """True when the document defines the named style.

    Must be checked *before* adding content: ``add_paragraph``/``add_heading``
    insert the paragraph first and only then resolve the style, so catching
    their ``KeyError`` afterwards would leave a stray paragraph behind.
    """
    try:
        document.styles[style_name]
    except KeyError:
        return False
    return True


def _add_heading_safely(document: Document, text: str, level: int):
    """Add a heading, falling back to a bold paragraph.

    Real-world documents do not always define the ``Heading N`` styles.
    """
    if _style_exists(document, f"Heading {level}"):
        return document.add_heading(text, level=level)
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True
    return paragraph


def _add_bullet_safely(document: Document, text: str):
    """Add a bulleted line, falling back to a plain ``- `` prefixed paragraph.

    The demo sample document defines no ``List Bullet`` style, so the style
    must not be assumed to exist.
    """
    if _style_exists(document, "List Bullet"):
        return document.add_paragraph(text, style="List Bullet")
    return document.add_paragraph(f"- {text}")


def _find_heading_paragraph(document: Document, heading_text: str):
    """Find the paragraph whose text equals the heading (case-insensitive).

    Plain text scan by design: routine pushes must not re-run the full
    feature-based parser.
    """
    wanted = (heading_text or "").strip().lower()
    if not wanted:
        return None
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() == wanted:
            return paragraph
    return None


def _build_update_block(
    document: Document,
    *,
    block_level: int,
    repository: str,
    branch: str,
    actor: str,
    before_sha: Optional[str],
    after_sha: str,
    project: ProjectConfig,
    changed_files: list[ChangedFile],
) -> list:
    """Append the marked update block at the end; return its paragraphs.

    The caller may relocate the returned paragraphs under a target heading.
    """
    paragraphs = []

    paragraphs.append(
        _add_heading_safely(document, AUTOMATED_SECTION_TITLE, level=block_level)
    )

    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    for label, value in (
        ("Timestamp", timestamp),
        ("Repository", repository),
        ("Branch", branch),
        ("Actor", actor),
        ("Before SHA", before_sha or "(none)"),
        ("After SHA", after_sha),
        ("Project ID", project.project_id),
        ("Document", project.document_name),
    ):
        paragraphs.append(document.add_paragraph(f"{label}: {value}"))

    paragraphs.append(
        _add_heading_safely(document, "Changed Files", level=min(block_level + 1, 9))
    )
    if changed_files:
        for changed in changed_files:
            if changed.old_path:
                text = f"{changed.change_type}: {changed.old_path} -> {changed.path}"
            else:
                text = f"{changed.change_type}: {changed.path}"
            paragraphs.append(_add_bullet_safely(document, text))
    else:
        paragraphs.append(
            document.add_paragraph(
                "No changed files were available for this run "
                "(first push or manual trigger)."
            )
        )

    paragraphs.append(document.add_paragraph(AUTOMATED_SECTION_NOTE))
    return paragraphs


def _relocate_after(anchor_paragraph, block_paragraphs) -> None:
    """Move the block (created at the document end) directly after the anchor."""
    for paragraph in reversed(block_paragraphs):
        anchor_paragraph._p.addnext(paragraph._p)


# ---------------------------------------------------------------------------
# main flow
# ---------------------------------------------------------------------------
def run_update(env: Mapping[str, str], repo_path: str = ".") -> UpdateResult:
    """Detect the push, route it via the skeleton, and update the DOCX."""
    warnings: list[str] = []

    repository = extract_repository_name(env.get("GITHUB_REPOSITORY", ""))
    running_in_ci = bool(env.get("GITHUB_SHA", "").strip())

    if running_in_ci:
        after_sha = env["GITHUB_SHA"].strip()
        branch = env.get("GITHUB_REF_NAME", "").strip() or "main"
        before_sha: Optional[str] = env.get("GITHUB_EVENT_BEFORE", "").strip() or None
    else:
        after_sha = "HEAD"
        branch = (
            _local_git_output(["rev-parse", "--abbrev-ref", "HEAD"], repo_path)
            or "main"
        )
        before_sha = (
            "HEAD~1"
            if _local_git_output(["rev-parse", "--verify", "HEAD~1"], repo_path)
            else None
        )

    actor = env.get("GITHUB_ACTOR", "").strip() or "local-user"

    project = resolve_project(repository)

    document_path = Path(repo_path) / project.document_location
    if not document_path.exists():
        raise FileNotFoundError(
            f"Configured demo document not found: {document_path} "
            f"(from document_location of {repository!r})"
        )

    # Changed files of this push.
    changed_files: list[ChangedFile] = []
    if is_missing_sha(before_sha):
        before_sha = None
        warnings.append(
            "Before SHA is missing or all zeroes; updating the document "
            "with an empty changed-file list."
        )
    else:
        try:
            change_set = build_change_set(
                repository=repository,
                branch=branch,
                before_sha=before_sha,
                after_sha=after_sha,
                repo_path=repo_path,
            )
            changed_files = change_set.changed_files
        except subprocess.CalledProcessError as error:
            warnings.append(
                f"git diff {before_sha}..{after_sha} failed "
                f"({error.stderr.strip() if error.stderr else error}); "
                "updating the document with an empty changed-file list."
            )

    # Skeleton: load the stored one; build it once if missing.
    skeleton_path = skeleton_path_for(project, repo_path)
    skeleton_created = False
    if skeleton_path.exists():
        skeleton = load_skeleton(skeleton_path)
    else:
        skeleton, skeleton_path = build_and_save_skeleton(project, repo_path)
        skeleton_created = True

    # Route the change (rule-based today, LLM later).
    change_summary = build_change_summary(changed_files)
    decision = route_change(change_summary, changed_files, skeleton)

    # Apply the update to the DOCX.
    document = Document(str(document_path))
    skeleton_updated = False

    if decision.decision == CREATE_NEW:
        _add_heading_safely(document, decision.new_heading, level=1)
        _build_update_block(
            document,
            block_level=2,
            repository=repository,
            branch=branch,
            actor=actor,
            before_sha=before_sha,
            after_sha=after_sha,
            project=project,
            changed_files=changed_files,
        )
        placement = f"new section {decision.new_heading!r} appended"

        append_section(skeleton, heading=decision.new_heading, level=1)
        save_skeleton(skeleton, skeleton_path)
        skeleton_updated = True
    else:
        anchor = _find_heading_paragraph(document, decision.target_heading)
        target_section = (
            find_section_by_id(skeleton, decision.target_section_id)
            if decision.target_section_id
            else None
        )
        block_level = min((target_section.level if target_section else 1) + 1, 9)

        block = _build_update_block(
            document,
            block_level=block_level,
            repository=repository,
            branch=branch,
            actor=actor,
            before_sha=before_sha,
            after_sha=after_sha,
            project=project,
            changed_files=changed_files,
        )

        if anchor is not None:
            _relocate_after(anchor, block)
            placement = f"under existing heading {decision.target_heading!r}"
        else:
            # Fallback: the block stays where it was built — the end.
            placement = "appended to end (target heading not found)"
            warnings.append(
                f"Target heading {decision.target_heading!r} was not found "
                "in the DOCX; appended the update to the end instead."
            )

    document.save(str(document_path))

    return UpdateResult(
        document_path=document_path,
        project_id=project.project_id,
        repository=repository,
        branch=branch,
        changed_files=changed_files,
        decision=decision,
        placement=placement,
        skeleton_path=skeleton_path,
        skeleton_created=skeleton_created,
        skeleton_updated=skeleton_updated,
        warnings=warnings,
    )


def format_result(result: UpdateResult) -> str:
    """Render a readable terminal confirmation of the update."""
    lines = [
        "=" * 60,
        "TechDocker Demo DOCX Updater",
        "=" * 60,
        f"Updated document:    {result.document_path}",
        f"Project ID:          {result.project_id}",
        f"Repository / branch: {result.repository} / {result.branch}",
        f"Changed files:       {len(result.changed_files)}",
        f"Routing decision:    {result.decision.decision if result.decision else '(none)'}",
        f"Placement:           {result.placement}",
        f"Skeleton:            {result.skeleton_path}"
        + (" (created)" if result.skeleton_created else "")
        + (" (updated)" if result.skeleton_updated else ""),
    ]
    if result.decision is not None:
        lines.append(f"Reasoning:           {result.decision.reasoning}")
    if result.warnings:
        lines += [f"! {warning}" for warning in result.warnings]
    lines.append("=" * 60)
    return "\n".join(lines)


def main() -> int:
    """Entry point for ``python3 -m src.demo_docx_updater``."""
    result = run_update(os.environ)
    print(format_result(result))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
