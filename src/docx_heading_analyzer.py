"""Orchestrate feature extraction and scoring for a whole DOCX document.

The analyzer is the single entry point used by both the DOCX parser and the
``analyze_headings`` inspection command. It:

1. extracts features for every non-empty paragraph,
2. detects the document's body formatting,
3. detects repeated non-body formatting patterns,
4. scores and classifies each paragraph,
5. refines predicted heading levels using font-size groups where numbering is
   absent,
6. can render the results as a pandas DataFrame.

basically, coordinates the analysis: it gets features, body formatting, 
repeated patterns, calls the scorer then 
converts results into records/DataFrames.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document

from src import heading_scorer
from src.docx_feature_extractor import (
    BodyFormatting,
    ParagraphFeatures,
    detect_body_formatting,
    extract_paragraph_features,
    read_document_defaults,
)
from src.docx_numbering import NumberingResolver
from src.format_pattern_detector import detect_repeated_patterns
from src.heading_scorer import ScoringResult


@dataclass
class AnalyzedParagraph:
    """A paragraph together with its extracted features and scoring result."""

    features: ParagraphFeatures
    result: ScoringResult


@dataclass
class DocumentAnalysis:
    """The full analysis of a document."""

    paragraphs: list[AnalyzedParagraph]
    body_formatting: BodyFormatting
    warnings: list[str] = field(default_factory=list)


def analyze_word_document(word_document: Any) -> DocumentAnalysis:
    """Analyze an already-loaded python-docx ``Document``."""
    numbering = NumberingResolver(word_document)
    defaults = read_document_defaults(word_document)

    features: list[ParagraphFeatures] = []
    for index, paragraph in enumerate(word_document.paragraphs, start=1):
        # Resolve list metadata for every paragraph (even empty ones) so the
        # sequential numbering counters track document order correctly.
        list_info = numbering.resolve(paragraph)
        if not paragraph.text.strip():
            continue
        features.append(
            extract_paragraph_features(paragraph, index, list_info, defaults)
        )

    body = detect_body_formatting(features)
    for feature in features:
        feature.body_font_size = body.font_size
        feature.body_font_family = body.font_family
        feature.body_font_color = body.font_color

    detect_repeated_patterns(features, body)

    analyzed = [
        AnalyzedParagraph(features=feature, result=heading_scorer.score_paragraph(feature))
        for feature in features
    ]

    _refine_predicted_levels(analyzed)

    return DocumentAnalysis(
        paragraphs=analyzed,
        body_formatting=body,
        warnings=list(numbering.warnings),
    )


def analyze_docx_file(file_path: str) -> DocumentAnalysis:
    """Load a DOCX file from disk and analyze it (read-only)."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Unsupported file type: {path.suffix}. Expected a .docx file.")
    return analyze_word_document(Document(path))


def _refine_predicted_levels(analyzed: list[AnalyzedParagraph]) -> None:
    """Infer levels for heuristic headings that lack a numbering prefix.

    Uses relative font-size groups within the document: the largest heading font
    becomes level 1, the next distinct size level 2, and so on (capped at 9).
    Paragraphs that still cannot be levelled keep ``predicted_level = None``.
    """
    candidates = [
        item
        for item in analyzed
        if item.result.detection_method == "formatting_heuristic"
        and item.result.classification in ("heading", "probable_heading")
    ]

    distinct_sizes = sorted(
        {
            item.features.font_size
            for item in candidates
            if item.features.font_size is not None
        },
        reverse=True,
    )
    size_rank = {size: rank + 1 for rank, size in enumerate(distinct_sizes)}

    for item in candidates:
        if item.result.predicted_level is not None:
            continue
        font_size = item.features.font_size
        if font_size is not None and font_size in size_rank:
            item.result.predicted_level = min(size_rank[font_size], 9)


# ---------------------------------------------------------------------------
# Pandas output
# ---------------------------------------------------------------------------
# Columns emitted by the inspection CSV, in order.
ANALYSIS_COLUMNS: list[str] = [
    "position",
    "text",
    "style",
    "score",
    "classification",
    "predicted_level",
    "detection_method",
    "signals",
    "numbering_prefix",
    "is_word_list_item",
    "list_type",
    "list_level",
    "numbering_id",
    "numbering_format",
    "list_marker",
    "display_text",
    "has_internal_line_breaks",
    "segment_count",
    "segments",
    "word_count",
    "sentence_count",
    "is_title_case",
    "is_all_caps",
    "font_size",
    "body_font_size",
    "font_family",
    "body_font_family",
    "font_color",
    "is_colored",
    "bold",
    "italic",
    "underlined",
    "spacing_before",
    "spacing_after",
    "alignment",
    "repeated_formatting_pattern",
]


def analysis_to_records(analysis: DocumentAnalysis) -> list[dict[str, Any]]:
    """Flatten an analysis into row dicts suitable for a DataFrame."""
    records: list[dict[str, Any]] = []
    for item in analysis.paragraphs:
        features = item.features
        result = item.result
        records.append(
            {
                "position": features.position,
                "text": features.text,
                "style": features.style_name,
                "score": result.score,
                "classification": result.classification,
                "predicted_level": result.predicted_level,
                "detection_method": result.detection_method,
                "signals": "; ".join(result.signals),
                "numbering_prefix": features.numbering_prefix,
                "is_word_list_item": features.is_word_list_item,
                "list_type": features.list_type,
                "list_level": features.list_level,
                "numbering_id": features.numbering_id,
                "numbering_format": features.numbering_format,
                "list_marker": features.list_marker,
                "display_text": features.display_text,
                "has_internal_line_breaks": features.has_internal_line_breaks,
                "segment_count": features.segment_count,
                "segments": " | ".join(features.segments),
                "word_count": features.word_count,
                "sentence_count": features.sentence_count,
                "is_title_case": features.is_title_case,
                "is_all_caps": features.is_all_caps,
                "font_size": features.font_size,
                "body_font_size": features.body_font_size,
                "font_family": features.font_family,
                "body_font_family": features.body_font_family,
                "font_color": features.font_color,
                "is_colored": features.is_colored,
                "bold": features.is_bold,
                "italic": features.is_italic,
                "underlined": features.is_underlined,
                "spacing_before": features.spacing_before,
                "spacing_after": features.spacing_after,
                "alignment": features.alignment,
                "repeated_formatting_pattern": features.repeated_formatting_pattern,
            }
        )
    return records


def analysis_to_dataframe(analysis: DocumentAnalysis):
    """Return a pandas DataFrame of the analysis (imported lazily)."""
    import pandas as pd

    return pd.DataFrame(analysis_to_records(analysis), columns=ANALYSIS_COLUMNS)
