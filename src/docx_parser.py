"""Parse a DOCX document into a structured heading hierarchy.

Two detection methods are used, in priority order:

1. **Official Word heading styles** (Heading 1-9) — always trusted, scored 100,
   level taken directly from the style.
2. **Formatting heuristic** — for paragraphs without an official style, a
   transparent 0-100 score decides whether they are treated as headings.

Heuristic paragraphs scoring 80+ with an inferable level are inserted into the
confirmed hierarchy. Probable headings (60-79) are recorded for review but never
silently confirmed. Note/Link lines and ordinary text remain content. The
original DOCX file is only read, never modified.
"""

from pathlib import Path
from typing import Any

from docx import Document

from src.docx_heading_analyzer import AnalyzedParagraph, analyze_word_document


def parse_docx_document(file_path: str) -> dict[str, Any]:
    """Read a DOCX document and convert it into a structured hierarchy.

    Preserves the existing output fields (``file_name``, ``file_type``,
    ``title``, ``headings``, ``unassigned_content``, ``tables``, ``warnings``)
    and adds ``probable_headings`` and ``analysis_summary``.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Unsupported file type: {path.suffix}. Expected a .docx file."
        )

    word_document = Document(path)
    analysis = analyze_word_document(word_document)

    parsed_document: dict[str, Any] = {
        "file_name": path.name,
        "file_type": "docx",
        "title": "",
        "headings": [],
        "unassigned_content": [],
        "tables": [],
        "probable_headings": [],
        "warnings": [],
        "analysis_summary": {
            "official_heading_count": 0,
            "heuristic_heading_count": 0,
            "probable_heading_count": 0,
            "normal_content_count": 0,
        },
    }

    heading_stack: list[dict[str, Any]] = []
    summary = parsed_document["analysis_summary"]

    for analyzed in analysis.paragraphs:
        features = analyzed.features
        result = analyzed.result
        text = features.text
        position = features.position
        style_name = features.style_name

        # The Word Title style becomes the document title, never a heading.
        if features.is_title_style and not parsed_document["title"]:
            parsed_document["title"] = text
            parsed_document["unassigned_content"].append(
                {
                    "id": f"paragraph-{position}",
                    "text": text,
                    "position": position,
                    "style": style_name,
                }
            )
            continue

        is_official = result.detection_method == "official_word_heading_style"
        is_confirmed_heuristic = (
            result.detection_method == "formatting_heuristic"
            and result.classification == "heading"
            and result.predicted_level is not None
        )

        if is_official or is_confirmed_heuristic:
            level = result.predicted_level
            heading = _build_heading(text, position, style_name, level, result)

            if is_official:
                summary["official_heading_count"] += 1
            else:
                summary["heuristic_heading_count"] += 1

            if level == 1 and not parsed_document["title"]:
                parsed_document["title"] = text

            while heading_stack and heading_stack[-1]["level"] >= level:
                heading_stack.pop()

            if heading_stack:
                heading_stack[-1]["children"].append(heading)
            else:
                parsed_document["headings"].append(heading)

            heading_stack.append(heading)
            continue

        # Everything else stays as content under the active heading.
        paragraph_data = {
            "id": f"paragraph-{position}",
            "text": text,
            "position": position,
            "style": style_name,
        }

        # One Word paragraph containing manual line breaks stays one paragraph
        # in the hierarchy, but its individual lines are exposed as segments.
        if features.has_internal_line_breaks:
            paragraph_data["has_internal_line_breaks"] = True
            paragraph_data["segments"] = features.segments
            paragraph_data["segment_count"] = features.segment_count

        # Word list items carry their resolved numbering metadata.
        if features.is_word_list_item:
            paragraph_data["is_word_list_item"] = True
            paragraph_data["list_type"] = features.list_type
            paragraph_data["list_level"] = features.list_level
            paragraph_data["numbering_id"] = features.numbering_id
            paragraph_data["numbering_format"] = features.numbering_format
            paragraph_data["list_marker"] = features.list_marker
            paragraph_data["display_text"] = features.display_text

        if result.classification == "probable_heading":
            summary["probable_heading_count"] += 1
            parsed_document["probable_headings"].append(
                {
                    "id": f"probable-{position}",
                    "text": text,
                    "position": position,
                    "style": style_name,
                    "score": result.score,
                    "predicted_level": result.predicted_level,
                    "detection_method": result.detection_method,
                    "signals": result.signals,
                }
            )
        elif (
            result.detection_method == "formatting_heuristic"
            and result.classification == "heading"
            and result.predicted_level is None
        ):
            # Strong heuristic heading but no reliable level — flag, do not
            # silently insert into the confirmed hierarchy.
            parsed_document["probable_headings"].append(
                {
                    "id": f"probable-{position}",
                    "text": text,
                    "position": position,
                    "style": style_name,
                    "score": result.score,
                    "predicted_level": None,
                    "detection_method": result.detection_method,
                    "signals": result.signals,
                }
            )
            parsed_document["warnings"].append(
                f"Paragraph at position {position} scored as a heading "
                f"({result.score}) but no reliable level could be inferred; "
                "flagged for review."
            )
        else:
            summary["normal_content_count"] += 1

        if heading_stack:
            heading_stack[-1]["content"].append(paragraph_data)
        else:
            parsed_document["unassigned_content"].append(paragraph_data)

            if not parsed_document["title"]:
                parsed_document["title"] = text

    parsed_document["tables"] = extract_tables(word_document)

    # Surface analyzer warnings (e.g. unresolvable list numbering definitions).
    parsed_document["warnings"].extend(analysis.warnings)

    if not parsed_document["headings"]:
        parsed_document["warnings"].append(
            "No confirmed headings were detected. "
            "The file may use manually formatted headings."
        )

    if parsed_document["probable_headings"]:
        parsed_document["warnings"].append(
            f"{len(parsed_document['probable_headings'])} probable heading(s) "
            "were detected and require manual review before use."
        )

    return parsed_document


def _build_heading(
    text: str,
    position: int,
    style_name: str,
    level: int,
    result: Any,
) -> dict[str, Any]:
    """Construct a confirmed-heading node for the hierarchy."""
    return {
        "id": f"heading-{position}",
        "title": text,
        "level": level,
        "position": position,
        "style": style_name,
        "content": [],
        "children": [],
        "detection_method": result.detection_method,
        "score": result.score,
        "signals": result.signals,
        "predicted_level": result.predicted_level,
    }


def extract_tables(word_document: Any) -> list[dict[str, Any]]:
    """Extract all tables into rows and cells."""
    extracted_tables: list[dict[str, Any]] = []

    for table_index, table in enumerate(word_document.tables, start=1):
        rows: list[list[str]] = []

        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells]
            rows.append(row_values)

        extracted_tables.append(
            {
                "id": f"table-{table_index}",
                "table_index": table_index,
                "rows": rows,
            }
        )

    return extracted_tables
