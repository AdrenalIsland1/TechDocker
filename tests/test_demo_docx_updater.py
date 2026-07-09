"""Tests for the demo DOCX updater (temporary DOCX files, no real git)."""

from __future__ import annotations

from docx import Document

from src import demo_docx_updater
from src.demo_docx_updater import AUTOMATED_SECTION_TITLE, run_update
from src.git_change_detector import ChangedFile, GitChangeSet
from src.project_resolver import ProjectConfig

CI_ENV = {
    "GITHUB_REPOSITORY": "AdrenalIsland1/TechDocker",
    "GITHUB_REF_NAME": "main",
    "GITHUB_SHA": "def456",
    "GITHUB_ACTOR": "Vaibhav",
    "GITHUB_EVENT_BEFORE": "abc123",
}


def make_demo_docx(tmp_path):
    """Create a small throwaway demo document."""
    path = tmp_path / "demo-document.docx"
    document = Document()
    document.add_heading("Existing Demo Document", level=1)
    document.add_paragraph("Original content that must survive the update.")
    document.save(path)
    return path


def make_project(document_path):
    return ProjectConfig(
        repository_name="TechDocker",
        project_id="techdocker",
        production_branch="main",
        document_name="demo-document.docx",
        document_location=str(document_path),
    )


def mock_resolver(monkeypatch, project):
    monkeypatch.setattr(
        demo_docx_updater, "resolve_project", lambda repository_name: project
    )


def mock_detector(monkeypatch, files):
    def fake_build_change_set(**kwargs):
        return GitChangeSet(
            repository=kwargs["repository"],
            branch=kwargs["branch"],
            before_sha=kwargs["before_sha"],
            after_sha=kwargs["after_sha"],
            changed_files=files,
        )

    monkeypatch.setattr(demo_docx_updater, "build_change_set", fake_build_change_set)


def docx_texts(path):
    return [paragraph.text for paragraph in Document(str(path)).paragraphs]


# ---------------------------------------------------------------------------
# tests
# ---------------------------------------------------------------------------
def test_updater_appends_update_heading(tmp_path, monkeypatch):
    document_path = make_demo_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(monkeypatch, [])

    run_update(CI_ENV, repo_path=str(tmp_path))

    document = Document(str(document_path))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    assert AUTOMATED_SECTION_TITLE in headings
    # Original content survives the append.
    assert "Original content that must survive the update." in docx_texts(
        document_path
    )


def test_changed_files_appear_in_docx_text(tmp_path, monkeypatch):
    document_path = make_demo_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(
        monkeypatch,
        [
            ChangedFile(path="src/example.py", change_type="modified"),
            ChangedFile(path="tests/test_example.py", change_type="added"),
        ],
    )

    result = run_update(CI_ENV, repo_path=str(tmp_path))

    texts = docx_texts(document_path)
    assert "modified: src/example.py" in texts
    assert "added: tests/test_example.py" in texts
    assert len(result.changed_files) == 2


def test_missing_before_sha_still_creates_update_section(tmp_path, monkeypatch):
    document_path = make_demo_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))

    def must_not_be_called(**kwargs):
        raise AssertionError("build_change_set must not run without a before SHA")

    monkeypatch.setattr(demo_docx_updater, "build_change_set", must_not_be_called)

    env = dict(CI_ENV, GITHUB_EVENT_BEFORE="0" * 40)
    result = run_update(env, repo_path=str(tmp_path))

    texts = docx_texts(document_path)
    assert result.changed_files == []
    assert any("No changed files were available" in text for text in texts)
    assert any("missing or all zeroes" in warning for warning in result.warnings)
    # The section itself still exists.
    assert AUTOMATED_SECTION_TITLE in [
        paragraph.text for paragraph in Document(str(document_path)).paragraphs
    ]


def test_project_and_document_metadata_appear_in_docx(tmp_path, monkeypatch):
    document_path = make_demo_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(monkeypatch, [])

    run_update(CI_ENV, repo_path=str(tmp_path))

    texts = docx_texts(document_path)
    assert "Project ID: techdocker" in texts
    assert "Document: demo-document.docx" in texts
    assert "Repository: TechDocker" in texts
    assert "Branch: main" in texts
    assert "Actor: Vaibhav" in texts
    assert "Before SHA: abc123" in texts
    assert "After SHA: def456" in texts
    assert any(
        "generated automatically by the TechDocker GitHub automation demo"
        in text
        for text in texts
    )


def test_document_without_list_bullet_style_falls_back_to_plain_bullets(
    tmp_path, monkeypatch
):
    # The real sample document defines no "List Bullet" style; the updater
    # must fall back to plain "- " paragraphs instead of raising KeyError.
    document_path = make_demo_docx(tmp_path)
    document = Document(str(document_path))
    styles = document.styles
    styles.element.remove(styles["List Bullet"].element)
    document.save(str(document_path))

    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(
        monkeypatch,
        [ChangedFile(path="src/example.py", change_type="modified")],
    )

    run_update(CI_ENV, repo_path=str(tmp_path))

    texts = docx_texts(document_path)
    assert "- modified: src/example.py" in texts
    # No stray duplicate from a partially-added styled paragraph.
    assert "modified: src/example.py" not in texts


def test_updater_returns_updated_path_and_prints_it(tmp_path, monkeypatch, capsys):
    document_path = make_demo_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(monkeypatch, [ChangedFile(path="a.py", change_type="added")])

    result = run_update(CI_ENV, repo_path=str(tmp_path))
    print(demo_docx_updater.format_result(result))

    assert result.document_path.name == "demo-document.docx"
    assert result.project_id == "techdocker"

    captured = capsys.readouterr().out
    assert "demo-document.docx" in captured
    assert "techdocker" in captured
    assert "Changed files:       1" in captured
