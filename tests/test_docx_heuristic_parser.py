"""Integration tests for heuristic DOCX heading detection.

These build real DOCX files with python-docx (including direct ``w:numPr``
numbering XML for list-item detection) and exercise both the analyzer and the
integrated ``parse_docx_document`` output.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.docx_heading_analyzer import analyze_word_document
from src.docx_parser import parse_docx_document


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def add_para(
    document,
    text,
    *,
    bold=False,
    size=None,
    color=None,
    space_before=None,
):
    """Add a paragraph with a single, optionally formatted run."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    return paragraph


def add_list_item(document, text, num_id=1):
    """Add a paragraph carrying a real Word ``w:numPr`` numbering property."""
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return paragraph


def analyze(document):
    """Return {text: AnalyzedParagraph} for convenience."""
    analysis = analyze_word_document(document)
    return {item.features.text: item for item in analysis.paragraphs}


# ---------------------------------------------------------------------------
# official styles still work
# ---------------------------------------------------------------------------
def test_official_heading_one_scores_100(tmp_path):
    document = Document()
    document.add_heading("Architecture", level=1)
    by_text = analyze(document)
    result = by_text["Architecture"].result
    assert result.score == 100
    assert result.detection_method == "official_word_heading_style"
    assert result.predicted_level == 1


def test_official_heading_two_keeps_level(tmp_path):
    document = Document()
    document.add_heading("Subsystem", level=2)
    result = analyze(document)["Subsystem"].result
    assert result.predicted_level == 2
    assert result.score == 100


# ---------------------------------------------------------------------------
# list-item vs numbering prefix
# ---------------------------------------------------------------------------
def test_real_word_list_item_gets_penalty():
    document = Document()
    item = add_list_item(document, "This is an ordinary bullet point item")
    by_text = analyze(document)
    analyzed = by_text["This is an ordinary bullet point item"]
    assert analyzed.features.is_word_list_item is True
    assert "word_list_item:-20" in analyzed.result.signals


def test_textual_numbering_prefix_gets_bonus():
    document = Document()
    add_para(document, "3.4.1 Authentication", bold=True, size=14)
    analyzed = analyze(document)["3.4.1 Authentication"]
    assert analyzed.features.numbering_prefix == "3.4.1"
    assert "numbering_prefix:+25" in analyzed.result.signals
    assert analyzed.result.predicted_level == 3


# ---------------------------------------------------------------------------
# colour
# ---------------------------------------------------------------------------
def test_colored_paragraph_is_penalized():
    document = Document()
    add_para(document, "Colored Heading Text", color=RGBColor(0xFF, 0x00, 0x00))
    analyzed = analyze(document)["Colored Heading Text"]
    assert analyzed.features.is_colored is True
    assert analyzed.features.font_color == "FF0000"
    assert "colored_text:-20" in analyzed.result.signals


def test_default_black_text_is_not_colored():
    document = Document()
    add_para(document, "Ordinary Heading Text")
    analyzed = analyze(document)["Ordinary Heading Text"]
    assert analyzed.features.is_colored is False
    assert "colored_text:-20" not in analyzed.result.signals


def test_explicit_black_text_is_not_colored():
    document = Document()
    add_para(document, "Explicit Black Text", color=RGBColor(0x00, 0x00, 0x00))
    analyzed = analyze(document)["Explicit Black Text"]
    assert analyzed.features.is_colored is False


# ---------------------------------------------------------------------------
# repeated formatting
# ---------------------------------------------------------------------------
def test_repeated_heading_formatting_gets_bonus():
    document = Document()
    add_para(document, "This is normal body text content one.", size=11)
    add_para(document, "This is normal body text content two.", size=11)
    add_para(document, "Overview", bold=True, size=16)
    add_para(document, "Details", bold=True, size=16)
    by_text = analyze(document)
    assert by_text["Overview"].features.repeated_formatting_pattern is True
    assert "repeated_pattern:+20" in by_text["Overview"].result.signals
    assert "repeated_pattern:+20" in by_text["Details"].result.signals


def test_repeated_body_formatting_gets_no_bonus():
    document = Document()
    add_para(document, "Normal body sentence number one here.", size=11)
    add_para(document, "Normal body sentence number two here.", size=11)
    add_para(document, "Normal body sentence number three here.", size=11)
    by_text = analyze(document)
    for text, item in by_text.items():
        assert item.features.repeated_formatting_pattern is False, text
        assert "repeated_pattern:+20" not in item.result.signals


def test_single_colored_line_does_not_poison_body_color_detection():
    # A single explicitly-coloured paragraph must not become the "body colour"
    # (which would strip the body-formatting penalty from ordinary sentences).
    document = Document()
    add_para(document, "Normal body sentence number one here.", size=11)
    add_para(document, "Normal body sentence number two here.", size=11)
    add_para(document, "A rare coloured line.", size=11, color=RGBColor(0xFF, 0, 0))
    by_text = analyze(document)
    body_line = by_text["Normal body sentence number one here."]
    assert body_line.features.body_font_color is None
    assert body_line.features.repeated_formatting_pattern is False
    assert "repeated_pattern:+20" not in body_line.result.signals


# ---------------------------------------------------------------------------
# integrated parser output
# ---------------------------------------------------------------------------
def test_strong_heuristic_heading_detected_without_official_styles(tmp_path):
    document = Document()
    add_para(document, "This system stores warehouse inventory records.", size=11)
    add_para(document, "1. System Architecture", bold=True, size=16)
    add_para(document, "The architecture is layered and modular.", size=11)
    file_path = tmp_path / "heuristic.docx"
    document.save(file_path)

    parsed = parse_docx_document(str(file_path))

    titles = [heading["title"] for heading in parsed["headings"]]
    assert "1. System Architecture" in titles
    heading = parsed["headings"][0]
    assert heading["level"] == 1
    assert heading["detection_method"] == "formatting_heuristic"
    assert heading["score"] >= 80
    assert parsed["analysis_summary"]["heuristic_heading_count"] == 1


def test_probable_heading_recorded_but_not_confirmed(tmp_path):
    document = Document()
    add_para(document, "This system stores warehouse inventory records.", size=11)
    add_para(document, "System Overview", bold=True, size=16, space_before=12)
    add_para(document, "It handles day to day warehouse activity.", size=11)
    file_path = tmp_path / "probable.docx"
    document.save(file_path)

    parsed = parse_docx_document(str(file_path))

    probable_texts = [item["text"] for item in parsed["probable_headings"]]
    assert "System Overview" in probable_texts

    confirmed_titles = [heading["title"] for heading in parsed["headings"]]
    assert "System Overview" not in confirmed_titles
    assert parsed["analysis_summary"]["probable_heading_count"] >= 1


def test_note_line_stays_normal_content_in_parser(tmp_path):
    document = Document()
    document.add_heading("Deployment", level=1)
    add_para(document, "Note: This method is deprecated.", bold=True)
    file_path = tmp_path / "note.docx"
    document.save(file_path)

    parsed = parse_docx_document(str(file_path))

    deployment = parsed["headings"][0]
    note_texts = [item["text"] for item in deployment["content"]]
    assert "Note: This method is deprecated." in note_texts
    probable_texts = [item["text"] for item in parsed["probable_headings"]]
    assert "Note: This method is deprecated." not in probable_texts


# ---------------------------------------------------------------------------
# robustness
# ---------------------------------------------------------------------------
def test_empty_document_does_not_crash(tmp_path):
    document = Document()
    file_path = tmp_path / "empty.docx"
    document.save(file_path)
    parsed = parse_docx_document(str(file_path))
    assert parsed["headings"] == []
    assert parsed["file_type"] == "docx"


def test_paragraph_with_no_explicit_formatting_does_not_crash(tmp_path):
    document = Document()
    document.add_paragraph("Plain paragraph with no explicit run formatting.")
    empty_run_para = document.add_paragraph()
    empty_run_para.add_run("")  # empty run, no visible text
    file_path = tmp_path / "plain.docx"
    document.save(file_path)
    parsed = parse_docx_document(str(file_path))
    assert parsed["file_type"] == "docx"
