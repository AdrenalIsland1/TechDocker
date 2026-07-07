"""Tests for manual line-break segmentation and Word list-marker resolution."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from src.docx_heading_analyzer import (
    analysis_to_dataframe,
    analyze_word_document,
)
from src.docx_parser import parse_docx_document


def analyze(document):
    analysis = analyze_word_document(document)
    return analysis, {item.features.text: item for item in analysis.paragraphs}


# ---------------------------------------------------------------------------
# Problem 1: manual line breaks inside one Word paragraph
# ---------------------------------------------------------------------------
def make_profile_paragraph(document):
    """One Word paragraph with manual line breaks between profile lines."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Name: Vaibhav")
    run.add_break()
    run.add_text("Age: 20")
    run.add_break()
    run.add_text("Height: 173 cm")
    run.add_break()
    run.add_text("Primary goal: run a marathon")
    return paragraph


def test_manual_line_breaks_produce_segments():
    document = Document()
    make_profile_paragraph(document)
    _, by_text = analyze(document)
    (features,) = [item.features for item in by_text.values()]

    assert features.has_internal_line_breaks is True
    assert features.segment_count == 4
    assert features.segments == [
        "Name: Vaibhav",
        "Age: 20",
        "Height: 173 cm",
        "Primary goal: run a marathon",
    ]


def test_full_text_with_newlines_is_preserved():
    document = Document()
    make_profile_paragraph(document)
    _, by_text = analyze(document)
    (features,) = [item.features for item in by_text.values()]

    assert "\n" in features.text
    assert features.text == (
        "Name: Vaibhav\nAge: 20\nHeight: 173 cm\nPrimary goal: run a marathon"
    )


def test_empty_segments_are_ignored():
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("First line")
    run.add_break()
    run.add_break()  # empty line between breaks
    run.add_text("Second line")
    run.add_break()  # trailing break
    _, by_text = analyze(document)
    (features,) = [item.features for item in by_text.values()]

    assert features.segments == ["First line", "Second line"]
    assert features.segment_count == 2


def test_paragraph_without_breaks_has_single_segment():
    document = Document()
    document.add_paragraph("Just one ordinary line.")
    _, by_text = analyze(document)
    features = by_text["Just one ordinary line."].features

    assert features.has_internal_line_breaks is False
    assert features.segment_count == 1
    assert features.segments == ["Just one ordinary line."]


def test_segmented_paragraph_stays_one_paragraph_in_hierarchy(tmp_path):
    document = Document()
    document.add_heading("Profile", level=1)
    make_profile_paragraph(document)
    file_path = tmp_path / "profile.docx"
    document.save(file_path)

    parsed = parse_docx_document(str(file_path))

    profile = parsed["headings"][0]
    assert len(profile["content"]) == 1  # still ONE paragraph, not four
    paragraph = profile["content"][0]
    assert paragraph["has_internal_line_breaks"] is True
    assert paragraph["segment_count"] == 4
    assert paragraph["segments"][0] == "Name: Vaibhav"


# ---------------------------------------------------------------------------
# Problem 2: Word list markers
# ---------------------------------------------------------------------------
def test_numbered_list_items_detected_with_sequential_markers():
    document = Document()
    texts = ["Improve running", "Sleep eight hours", "Drink more water"]
    for text in texts:
        document.add_paragraph(text, style="List Number")
    _, by_text = analyze(document)

    for expected_marker, text in zip(["1.", "2.", "3."], texts):
        features = by_text[text].features
        assert features.is_word_list_item is True
        assert features.list_type == "numbered"
        assert features.numbering_format == "decimal"
        assert features.list_marker == expected_marker
        assert features.display_text == f"{expected_marker} {text}"


def test_list_order_matches_document_order():
    document = Document()
    document.add_paragraph("Alpha", style="List Number")
    document.add_paragraph("An ordinary body paragraph in between.")
    document.add_paragraph("Beta", style="List Number")
    _, by_text = analyze(document)

    assert by_text["Alpha"].features.list_marker == "1."
    assert by_text["Beta"].features.list_marker == "2."
    assert by_text["An ordinary body paragraph in between."].features.is_word_list_item is False


def test_bullet_list_items():
    document = Document()
    document.add_paragraph("Stretching", style="List Bullet")
    document.add_paragraph("Hydration", style="List Bullet")
    _, by_text = analyze(document)

    for text in ("Stretching", "Hydration"):
        features = by_text[text].features
        assert features.is_word_list_item is True
        assert features.list_type == "bullet"
        assert features.numbering_format == "bullet"
        assert features.list_marker == "•"
        assert features.display_text == f"• {text}"


def _add_multilevel_numbering(document, num_id=99):
    """Inject a two-level decimal numbering definition into numbering.xml."""
    root = document.part.part_related_by(RT.NUMBERING).element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(num_id))
    for ilvl, lvl_text in ((0, "%1."), (1, "%1.%2.")):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        text_el = OxmlElement("w:lvlText")
        text_el.set(qn("w:val"), lvl_text)
        lvl.append(start)
        lvl.append(num_fmt)
        lvl.append(text_el)
        abstract.append(lvl)

    first_num = root.find(qn("w:num"))
    root.insert(list(root).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(num_id))
    num.append(abstract_ref)
    root.append(num)


def _add_numbered_item(document, text, num_id, ilvl):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return paragraph


def test_nested_list_levels_and_markers():
    document = Document()
    _add_multilevel_numbering(document, num_id=99)
    _add_numbered_item(document, "Top one", 99, 0)
    _add_numbered_item(document, "Sub one", 99, 1)
    _add_numbered_item(document, "Sub two", 99, 1)
    _add_numbered_item(document, "Top two", 99, 0)
    _add_numbered_item(document, "Fresh sub", 99, 1)  # counter must restart
    _, by_text = analyze(document)

    assert by_text["Top one"].features.list_level == 0
    assert by_text["Sub one"].features.list_level == 1
    assert by_text["Top one"].features.list_marker == "1."
    assert by_text["Sub one"].features.list_marker == "1.1."
    assert by_text["Sub two"].features.list_marker == "1.2."
    assert by_text["Top two"].features.list_marker == "2."
    assert by_text["Fresh sub"].features.list_marker == "2.1."


def test_unresolvable_numbering_gets_no_marker_and_warns():
    document = Document()
    _add_numbered_item(document, "Orphan item", num_id=777, ilvl=0)
    analysis, by_text = analyze(document)

    features = by_text["Orphan item"].features
    assert features.is_word_list_item is True
    assert features.list_type == "unknown"
    assert features.list_marker is None
    assert features.numbering_id == 777
    assert any("777" in warning for warning in analysis.warnings)


def test_word_numbering_is_not_a_textual_numbering_prefix():
    document = Document()
    document.add_paragraph("Improve running", style="List Number")
    _, by_text = analyze(document)

    analyzed = by_text["Improve running"]
    # The visible "1." lives in XML, not the text: no +25 prefix bonus...
    assert analyzed.features.numbering_prefix is None
    assert "numbering_prefix:+25" not in analyzed.result.signals
    # ...but the real list item still receives the -20 penalty.
    assert "word_list_item:-20" in analyzed.result.signals


def test_manual_textual_numbering_still_gets_prefix_bonus():
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("2.1 API Configuration:")
    run.bold = True
    _, by_text = analyze(document)

    analyzed = by_text["2.1 API Configuration:"]
    assert analyzed.features.numbering_prefix == "2.1"
    assert analyzed.features.is_word_list_item is False
    assert "numbering_prefix:+25" in analyzed.result.signals
    assert "word_list_item:-20" not in analyzed.result.signals


# ---------------------------------------------------------------------------
# body detection in list-heavy documents
# ---------------------------------------------------------------------------
def test_list_items_count_towards_body_formatting():
    """In a list-heavy document, bold pseudo-headings must not become "body".

    The default font (from w:docDefaults) plus the list items should win the
    body vote, so the bold larger titles keep their heading-like signals and
    the list items receive the body-formatting penalty — with no false
    repeated-pattern bonus for ordinary content.
    """
    from docx.shared import Pt

    document = Document()
    for title in ("Personal Details", "Priorities", "Habits"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
    for text in ("Improve running", "Sleep eight hours", "Drink more water"):
        document.add_paragraph(text, style="List Number")
    _, by_text = analyze(document)

    title_item = by_text["Personal Details"]
    assert title_item.features.body_font_size is not None
    assert title_item.features.body_font_size < 14
    assert "larger_font:+20" in title_item.result.signals

    list_item = by_text["Improve running"]
    assert "body_formatting:-15" in list_item.result.signals
    assert "repeated_pattern:+20" not in list_item.result.signals


# ---------------------------------------------------------------------------
# CSV / report integration
# ---------------------------------------------------------------------------
def test_csv_contains_segment_and_list_columns():
    document = Document()
    make_profile_paragraph(document)
    document.add_paragraph("Improve running", style="List Number")
    analysis, _ = analyze(document)
    dataframe = analysis_to_dataframe(analysis)

    for column in (
        "has_internal_line_breaks",
        "segment_count",
        "segments",
        "list_type",
        "list_level",
        "numbering_id",
        "numbering_format",
        "list_marker",
        "display_text",
    ):
        assert column in dataframe.columns

    profile_row = dataframe[dataframe["has_internal_line_breaks"]].iloc[0]
    assert profile_row["segment_count"] == 4
    assert profile_row["segments"] == (
        "Name: Vaibhav | Age: 20 | Height: 173 cm | Primary goal: run a marathon"
    )

    list_row = dataframe[dataframe["is_word_list_item"]].iloc[0]
    assert list_row["list_marker"] == "1."
    assert list_row["display_text"] == "1. Improve running"
