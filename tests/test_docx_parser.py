from pathlib import Path

from docx import Document

from src.docx_parser import parse_docx_document


def test_docx_parser_reads_heading_hierarchy(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "project-master.docx"

    document = Document()

    document.add_heading(
        "Customer Platform",
        level=0,
    )

    document.add_heading(
        "Architecture",
        level=1,
    )

    document.add_heading(
        "Customer Service",
        level=2,
    )

    document.add_paragraph(
        "Handles customer operations."
    )

    document.save(file_path)

    result = parse_docx_document(str(file_path))

    assert result["file_type"] == "docx"
    assert result["title"] == "Customer Platform"

    architecture = result["headings"][0]

    assert architecture["title"] == "Architecture"
    assert architecture["level"] == 1

    customer_service = architecture["children"][0]

    assert customer_service["title"] == "Customer Service"
    assert customer_service["level"] == 2

    assert customer_service["content"][0]["text"] == (
        "Handles customer operations."
    )


def test_docx_parser_reads_tables(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "api-document.docx"

    document = Document()

    document.add_heading(
        "API Document",
        level=0,
    )

    table = document.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Method"
    table.cell(0, 1).text = "Path"
    table.cell(1, 0).text = "GET"
    table.cell(1, 1).text = "/health"

    document.save(file_path)

    result = parse_docx_document(str(file_path))

    assert len(result["tables"]) == 1

    assert result["tables"][0]["rows"] == [
        ["Method", "Path"],
        ["GET", "/health"],
    ]


def test_docx_parser_warns_when_no_headings_exist(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "unstructured.docx"

    document = Document()

    document.add_paragraph(
        "This document has no heading styles."
    )

    document.add_paragraph(
        "Everything is formatted as normal text."
    )

    document.save(file_path)

    result = parse_docx_document(str(file_path))

    assert result["headings"] == []

    assert len(result["warnings"]) == 1