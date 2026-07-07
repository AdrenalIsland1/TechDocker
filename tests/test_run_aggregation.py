"""Tests for visible-character-weighted run aggregation (bold/italic/underline)."""

from __future__ import annotations

from docx import Document

from src.docx_feature_extractor import extract_paragraph_features

LONG = "This is a long stretch of text that dominates the paragraph"
SHORT = "tiny"


def features_for(paragraph):
    return extract_paragraph_features(paragraph, position=1)


def two_run_paragraph(long_format: dict, short_format: dict):
    """One paragraph: a long run and a short run with the given formatting."""
    document = Document()
    paragraph = document.add_paragraph()
    long_run = paragraph.add_run(LONG)
    short_run = paragraph.add_run(" " + SHORT)
    for attribute, value in long_format.items():
        setattr(long_run.font, attribute, value)
    for attribute, value in short_format.items():
        setattr(short_run.font, attribute, value)
    return paragraph


# ---------------------------------------------------------------------------
# bold
# ---------------------------------------------------------------------------
def test_long_bold_run_with_short_plain_run_is_mostly_bold():
    paragraph = two_run_paragraph({"bold": True}, {"bold": False})
    assert features_for(paragraph).is_bold is True


def test_short_bold_run_with_long_plain_run_is_not_mostly_bold():
    paragraph = two_run_paragraph({"bold": False}, {"bold": True})
    assert features_for(paragraph).is_bold is False


# ---------------------------------------------------------------------------
# underline
# ---------------------------------------------------------------------------
def test_long_underlined_run_with_short_plain_run_is_mostly_underlined():
    paragraph = two_run_paragraph({"underline": True}, {"underline": False})
    assert features_for(paragraph).is_underlined is True


def test_short_underlined_run_with_long_plain_run_is_not_mostly_underlined():
    paragraph = two_run_paragraph({"underline": False}, {"underline": True})
    assert features_for(paragraph).is_underlined is False


# ---------------------------------------------------------------------------
# italic (50% threshold)
# ---------------------------------------------------------------------------
def test_long_italic_run_crosses_half_threshold():
    paragraph = two_run_paragraph({"italic": True}, {"italic": False})
    assert features_for(paragraph).is_italic is True


def test_short_italic_run_stays_below_half_threshold():
    paragraph = two_run_paragraph({"italic": False}, {"italic": True})
    assert features_for(paragraph).is_italic is False


# ---------------------------------------------------------------------------
# whitespace-only and empty runs
# ---------------------------------------------------------------------------
def test_whitespace_runs_do_not_affect_the_result():
    document = Document()
    paragraph = document.add_paragraph()
    bold_run = paragraph.add_run("Almost entirely bold text here")
    bold_run.font.bold = True
    for _ in range(4):  # several whitespace/empty runs, some marked non-bold
        filler = paragraph.add_run("   ")
        filler.font.bold = False
        paragraph.add_run("")

    features = features_for(paragraph)
    assert features.is_bold is True


def test_paragraph_with_only_whitespace_runs_does_not_crash():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("   ")
    paragraph.add_run(" ")

    features = features_for(paragraph)
    assert features.is_bold is False
    assert features.is_underlined is False
    assert features.is_italic is False
