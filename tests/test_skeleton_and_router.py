"""Tests for the skeleton builder, skeleton store, and change router."""

from __future__ import annotations

import json

from docx import Document

from src import demo_docx_updater
from src.change_router import CREATE_NEW, UPDATE_EXISTING, route_change
from src.demo_docx_updater import run_update
from src.document_skeleton_builder import build_and_save_skeleton, build_skeleton
from src.git_change_detector import ChangedFile, GitChangeSet
from src.project_resolver import ProjectConfig
from src.skeleton_store import (
    DocumentSkeleton,
    append_section,
    find_section_by_heading,
    find_section_by_id,
    load_skeleton,
    save_skeleton,
)

CI_ENV = {
    "GITHUB_REPOSITORY": "AdrenalIsland1/TechDocker",
    "GITHUB_REF_NAME": "main",
    "GITHUB_SHA": "def456",
    "GITHUB_ACTOR": "Vaibhav",
    "GITHUB_EVENT_BEFORE": "abc123",
}


def make_structured_docx(tmp_path, name="structured.docx"):
    """Temp DOCX with a small real heading hierarchy."""
    path = tmp_path / name
    document = Document()
    document.add_heading("System Overview", level=1)
    document.add_paragraph("The system parses documents.")
    document.add_heading("API Configuration", level=2)
    document.add_paragraph("Configure the API endpoint here.")
    document.add_heading("Validation Rules", level=1)
    document.add_paragraph("All inputs are validated.")
    document.save(path)
    return path


def make_project(document_path):
    return ProjectConfig(
        repository_name="TechDocker",
        project_id="techdocker",
        production_branch="main",
        document_name="structured.docx",
        document_location=str(document_path),
    )


def make_skeleton(headings):
    """In-memory skeleton with the given (heading, level) entries."""
    skeleton = DocumentSkeleton(
        source_document="test.docx", project_id="techdocker", generated_at="now"
    )
    for heading, level in headings:
        append_section(skeleton, heading=heading, level=level)
    return skeleton


def mock_resolver(monkeypatch, project):
    monkeypatch.setattr(
        demo_docx_updater, "resolve_project", lambda repository_name: project
    )


def mock_detector(monkeypatch, files):
    def fake_build_change_set(**kwargs):
        return GitChangeSet(
            repository=kwargs["repository"],
            branch=kwargs["branch"],
            before_sha=kwargs["before_sha"],
            after_sha=kwargs["after_sha"],
            changed_files=files,
        )

    monkeypatch.setattr(demo_docx_updater, "build_change_set", fake_build_change_set)


# ---------------------------------------------------------------------------
# skeleton builder
# ---------------------------------------------------------------------------
def test_skeleton_generation_creates_json_with_sections(tmp_path):
    document_path = make_structured_docx(tmp_path)
    project = make_project(document_path)

    skeleton, path = build_and_save_skeleton(project, repo_path=str(tmp_path))

    assert path.exists()
    data = json.loads(path.read_text(encoding="utf-8"))
    assert data["project_id"] == "techdocker"
    assert len(data["sections"]) == 3

    headings = [section["heading"] for section in data["sections"]]
    assert headings == ["System Overview", "API Configuration", "Validation Rules"]

    api_section = data["sections"][1]
    assert api_section["section_id"] == "system-overview-api-configuration"
    assert api_section["parent_id"] == "system-overview"
    assert api_section["path"] == "System Overview > API Configuration"
    assert api_section["level"] == 2
    assert api_section["order"] == 2
    assert api_section["content_hash"]  # section has content


def test_skeleton_ids_are_stable_across_rebuilds(tmp_path):
    document_path = make_structured_docx(tmp_path)
    first = build_skeleton(document_path, "techdocker", "structured.docx")
    second = build_skeleton(document_path, "techdocker", "structured.docx")
    assert [s.section_id for s in first.sections] == [
        s.section_id for s in second.sections
    ]


# ---------------------------------------------------------------------------
# skeleton store
# ---------------------------------------------------------------------------
def test_skeleton_save_and_load_roundtrip(tmp_path):
    skeleton = make_skeleton([("System Overview", 1), ("Validation Rules", 1)])
    path = tmp_path / "skeleton.json"

    save_skeleton(skeleton, path)
    loaded = load_skeleton(path)

    assert loaded.project_id == "techdocker"
    assert [s.heading for s in loaded.sections] == [
        "System Overview",
        "Validation Rules",
    ]
    assert find_section_by_id(loaded, "system-overview").heading == "System Overview"
    assert find_section_by_heading(loaded, "validation rules") is not None

    new_section = append_section(loaded, "Deployment", level=1)
    assert new_section.section_id == "deployment"
    assert new_section.order == 3


# ---------------------------------------------------------------------------
# change router
# ---------------------------------------------------------------------------
def test_router_chooses_validation_rules_for_test_files():
    skeleton = make_skeleton([("System Overview", 1), ("Validation Rules", 1)])
    decision = route_change(
        "tests changed",
        [ChangedFile(path="tests/test_parser.py", change_type="modified")],
        skeleton,
    )
    assert decision.decision == UPDATE_EXISTING
    assert decision.target_heading == "Validation Rules"


def test_router_creates_section_when_category_heading_missing():
    skeleton = make_skeleton([("System Overview", 1)])
    decision = route_change(
        "tests changed",
        [ChangedFile(path="tests/test_parser.py", change_type="modified")],
        skeleton,
    )
    assert decision.decision == CREATE_NEW
    assert decision.new_heading == "Validation Rules"


def test_router_falls_back_safely():
    # No keyword match: System Overview preferred, else the first heading.
    skeleton = make_skeleton([("Introduction", 1), ("Appendix", 1)])
    decision = route_change(
        "misc change",
        [ChangedFile(path="src/misc.py", change_type="modified")],
        skeleton,
    )
    assert decision.decision == UPDATE_EXISTING
    assert decision.target_heading == "Introduction"

    # Empty skeleton: create the fallback section rather than crash.
    empty = DocumentSkeleton(
        source_document="x.docx", project_id="techdocker", generated_at="now"
    )
    decision = route_change("misc change", [], empty)
    assert decision.decision == CREATE_NEW
    assert decision.new_heading == "System Overview"


# ---------------------------------------------------------------------------
# probable headings in the skeleton
# ---------------------------------------------------------------------------
def make_probable_docx(tmp_path, name="probable.docx"):
    """Temp DOCX whose 'Validation Rules:' is a probable heading (score 60-79).

    Three short bold paragraphs share one formatting signature (repeated
    pattern); the colon ending pushes only 'Validation Rules:' to 65:
    bold +15, few_words +10, repeated_pattern +20, title_or_caps +5,
    no_full_stop +5, colon_ending +10.
    """
    path = tmp_path / name
    document = Document()
    document.add_heading("System Overview", level=1)
    document.add_paragraph("The system parses documents and scores headings.")
    for text in ("Deployment Steps", "Final Notes", "Validation Rules:"):
        paragraph = document.add_paragraph()
        paragraph.add_run(text).bold = True
    document.add_paragraph("All inputs are validated before processing.")
    document.save(path)
    return path


def test_probable_headings_are_included_in_skeleton(tmp_path):
    document_path = make_probable_docx(tmp_path)

    skeleton = build_skeleton(document_path, "techdocker", "probable.docx")

    validation = find_section_by_heading(skeleton, "Validation Rules")
    assert validation is not None
    assert validation.heading == "Validation Rules:"
    assert validation.classification == "probable_heading"
    assert validation.score is not None and 60 <= validation.score < 80

    overview = find_section_by_heading(skeleton, "System Overview")
    assert overview.classification == "heading"
    assert overview.score == 100  # official Word heading style


def test_find_section_by_heading_ignores_trailing_colon():
    skeleton = make_skeleton([("Validation Rules:", 1)])
    assert find_section_by_heading(skeleton, "Validation Rules") is not None
    assert find_section_by_heading(skeleton, "validation rules:-") is not None
    assert find_section_by_heading(skeleton, "Validation") is None


def test_test_file_change_routes_to_existing_probable_section(tmp_path):
    document_path = make_probable_docx(tmp_path)
    skeleton = build_skeleton(document_path, "techdocker", "probable.docx")

    decision = route_change(
        "tests changed",
        [ChangedFile(path="tests/test_parser.py", change_type="modified")],
        skeleton,
    )

    assert decision.decision == UPDATE_EXISTING
    assert decision.target_heading == "Validation Rules:"


def test_updater_places_block_under_existing_probable_section(
    tmp_path, monkeypatch
):
    document_path = make_probable_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(
        monkeypatch,
        [ChangedFile(path="tests/test_parser.py", change_type="modified")],
    )

    result = run_update(CI_ENV, repo_path=str(tmp_path))

    assert result.decision.decision == UPDATE_EXISTING
    assert "Validation Rules:" in result.placement
    assert result.skeleton_updated is False  # no duplicate section created

    texts = [p.text for p in Document(str(document_path)).paragraphs]
    # No duplicate "Validation Rules" heading was appended...
    validation_headings = [
        t for t in texts if t.strip().lower().startswith("validation rules")
    ]
    assert validation_headings == ["Validation Rules:"]
    # ...and the block sits directly under the existing one.
    anchor = texts.index("Validation Rules:")
    assert texts[anchor + 1] == "Automated Documentation Update"


# ---------------------------------------------------------------------------
# updater integration with the skeleton
# ---------------------------------------------------------------------------
def test_updater_creates_skeleton_if_missing(tmp_path, monkeypatch):
    document_path = make_structured_docx(tmp_path)
    mock_resolver(monkeypatch, make_project(document_path))
    mock_detector(
        monkeypatch, [ChangedFile(path="src/misc.py", change_type="modified")]
    )

    result = run_update(CI_ENV, repo_path=str(tmp_path))

    assert result.skeleton_created is True
    assert result.skeleton_path.exists()
    data = json.loads(result.skeleton_path.read_text(encoding="utf-8"))
    assert len(data["sections"]) == 3
    # Fallback routing targeted the existing System Overview section.
    assert result.decision.target_heading == "System Overview"
    assert "under existing heading" in result.placement


def test_updater_updates_skeleton_when_new_section_created(tmp_path, monkeypatch):
    document_path = make_structured_docx(tmp_path)
    project = make_project(document_path)
    mock_resolver(monkeypatch, project)
    mock_detector(
        monkeypatch,
        [ChangedFile(path="src/docx_widget.py", change_type="added")],
    )

    # Pre-build the skeleton; it has no parser/document-processing section.
    build_and_save_skeleton(project, repo_path=str(tmp_path))

    result = run_update(CI_ENV, repo_path=str(tmp_path))

    assert result.decision.decision == CREATE_NEW
    assert result.decision.new_heading == "Document Processing"
    assert result.skeleton_created is False
    assert result.skeleton_updated is True

    reloaded = load_skeleton(result.skeleton_path)
    assert find_section_by_heading(reloaded, "Document Processing") is not None

    # The DOCX got the new heading too.
    texts = [p.text for p in Document(str(document_path)).paragraphs]
    assert "Document Processing" in texts
